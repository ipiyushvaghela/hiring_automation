from docx import Document
from io import BytesIO
import zipfile
import os
from config import TEXT_EXTRACTION_API_URL, TEST_API_URL
import requests
import uuid
from resume_parser import extract_text, pdf_to_base64
from resume_db import fetch_candidate_email
import pandas as pd
import csv
from flask import Flask, jsonify

app = Flask(__name__)

UPLOAD_DIRECTORY = "uploaded_data" # Directory to store uploaded files
EXTRACT_FOLDER = 'uploaded_data/extracted_files'


def create_docx(content):
    doc = Document()
    doc.add_heading('Generated Job Description', 0)
    doc.add_paragraph(content)
    
    # Save to BytesIO object
    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return byte_io

def save_uploaded_file(uploaded_file):
    file_path = os.path.join(UPLOAD_DIRECTORY, uploaded_file.name)
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    return file_path


def extract_zip_file(zip_file_path):
    # Ensure the extraction directory exists
    os.makedirs(EXTRACT_FOLDER, exist_ok=True)
  
    # Extract the zip file to the unique directory
    with zipfile.ZipFile(zip_file_path, 'r') as zip_ref:
        zip_ref.extractall(EXTRACT_FOLDER)
        print(f"Extracted files to {EXTRACT_FOLDER}")

    # Optionally remove the zip file after extraction
    if os.path.exists(zip_file_path):
        os.remove(zip_file_path)
        print(f"Removed zip file: {zip_file_path}")
    extract_dir = os.path.join(EXTRACT_FOLDER, os.path.basename(zip_file_path).split(".zip")[0])
    
    # Return the path of the directory where files were extracted
    return extract_dir
        
def process_resumes(base64_string, job_id):
    
    headers = {
        'Content-Type': 'application/json'  # Set the content type to JSON
    }
    response = requests.post(TEXT_EXTRACTION_API_URL, 
                             json={'file': base64_string, 'job_id': job_id},
                             headers=headers)
    
    print(response.status_code)
    if response.status_code == 200:
        print("Successfully processed resume")
    else:
        print("Error processing resume")
        

def get_candidates(candidate_data):
    return jsonify(candidate_data)

    # headers = {
    #     'Content-Type': 'application/json'  # Set the content type to JSON
    # }
    # response = requests.post(TEST_API_URL, 
    #                          json={'candidate': candidate_data},
    #                          headers=headers)
    
    # print(response.status_code)
    # if response.status_code == 200:
    #     print("Successfully processed resume")
    # else:
    #     print("Error processing resume")
        
def send_profile(profiles):
    if profiles:
        job_title = profiles[0][5].strip()
        print(f"Job Title: {job_title}")
        csv_file_name = f"candidate_profiles_{job_title}.csv"
        headers = ["ID", "Name", "Email", "Phone", "quesionnaire_response" ,"Job Title", "Job Id", "Rating"]

        # Create and save the CSV file
        with open(csv_file_name, mode='w', newline='') as file:
            writer = csv.writer(file)
            writer.writerow(headers)
            writer.writerows(profiles)

        print(f"CSV file generated: {csv_file_name}")
        return csv_file_name

