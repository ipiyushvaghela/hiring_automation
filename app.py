import streamlit as st
from config import ip_config,config,environment
import requests
import tempfile
from prompts import jd_prompt
from docx import Document
from io import BytesIO
from generate_questions import get_completion

upload_url = f"http://{ip_config[environment]}:{config['ResumeApplication']['Backend']}/upload"
fetch_url = f"http://{ip_config[environment]}:{config['ResumeApplication']['Backend']}/send_details"
fetch_email = f"http://{ip_config[environment]}:{config['ResumeApplication']['Backend']}/fetch_email"
fetch_questions = f"http://{ip_config[environment]}:{config['ResumeApplication']['Backend']}/fetch_questions"



def parse_resume(temp_file_path):
    """Send resume file path to backend for parsing."""
    response = requests.post(upload_url, json={"file": temp_file_path})
    if response.status_code == 200:
        st.success("Resume Parsed successfully")
    else:
        st.error("Error in parsing resume")

def save_temp_file(uploaded_file):
    """Save uploaded file to a temporary location."""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
        temp_file.write(uploaded_file.getvalue())
        return temp_file.name
    
def get_file_name_from_response():
    """Retrieve file name from backend response."""
    response = requests.post(fetch_url, json={"file": get_temp_file_name()})
    return response.json()["file_name"]

def fetch_and_display_questions(job_description, designation, email_id):
    """Fetch and display questions based on job description and email ID."""
    questions = {"job_description": job_description, "designation": designation, "email": email_id}
    response = requests.post(fetch_questions, json=questions)
    st.write(response.json())

def fetch_email_id():
    """Fetch email ID from backend."""
    file_name = get_file_name_from_response()
    response = requests.post(fetch_url, json={"file": file_name})
    email_data = {"data": response.json().get("data")}
    response = requests.post(fetch_email, json=email_data)
    return response.json()["email"]
    
def upload_resume():
    """Handle resume upload and parsing."""
    if st.button("Upload Resume"):
        uploaded_file = st.file_uploader("Choose a file")
        if uploaded_file is not None:
            job_description = st.text_area("Enter Job Description")
            designation = st.text_input("Enter Designation")
            if st.button("Submit"):
                temp_file_path = save_temp_file(uploaded_file)
                parse_resume(temp_file_path)
                email_id = fetch_email_id()
                fetch_and_display_questions(job_description, designation, email_id)

        
def jd_generator():
    if 'show_sidebar' not in st.session_state:
        st.session_state.show_sidebar = False

    if st.button("JD Generator"):
        st.session_state.show_sidebar = not st.session_state.show_sidebar

    if st.session_state.show_sidebar:
        render_sidebar()

def create_docx(content):
    doc = Document()
    doc.add_heading('Generated Job Description', 0)
    doc.add_paragraph(content)
    
    # Save to BytesIO object
    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return byte_io

def render_sidebar():
    """Render sidebar for filter selection."""
    # designation = st.sidebar.text_input("Designation")
    designation = "Senior Data Scientist"
    
    experirence = st.sidebar.text_input("Experience")
    skills = st.sidebar.text_area("Skills")
    technologies = st.sidebar.text_area("Technologies")
    other = st.sidebar.text_area("Other Details")

    if st.sidebar.button("Submit"):
        # data = {"experience":experirence,"skills":skills,"technologies":technologies,"other":other}
        data = {"experience":"5 years","skills":"python, tableu, my sql","technologies":"Genai","other":"should be good at communication"}
        
        jd_prompt[0]["content"] = jd_prompt[0]["content"].replace("<FILTERS>",str(data)).replace("<DESIGNATION>",designation)
        response = get_completion(jd_prompt)
        # Center the response text
        st.markdown("<div style='text-align: center;'>", unsafe_allow_html=True)
        st.markdown(f"**Generated Job Description:**\n\n{response}")
        st.markdown("</div>", unsafe_allow_html=True)
        st.write("<div style='text-align: center;'><h3>Generated Job Description</h3></div>", unsafe_allow_html=True)
        st.markdown(f"<div style='text-align: center;'>{response}</div>", unsafe_allow_html=True)
        # Create a .docx file for download
        docx_file = create_docx(response)
        
        st.download_button(
            label="Download as DOCX",
            data=docx_file,
            file_name="job_description.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
      
      
def main():
    st.markdown("""
        <style>
        .title {
            text-align: center;
        }
        </style>
    """, unsafe_allow_html=True)

    # Centered title
    st.markdown('<h1 class="title">Candidate Profile Solution</h1>', unsafe_allow_html=True)
    st.markdown("<br>", unsafe_allow_html=True)

    # Create empty columns for centering
    col_empty1, col1, col_mid, col2, col_empty2 = st.columns([1, 2, 0.5, 2, 1])

    with col1:
        upload_resume()
    with col2:
        jd_generator()

if __name__ == '__main__':
    main()
      
        
        

    